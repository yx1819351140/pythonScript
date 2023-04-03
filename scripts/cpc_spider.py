# -*- coding: UTF-8 -*-
# @Project -> File     ：pythonScript -> cpc_spider         
# @IDE      ：PyCharm
# @Author   ：yangxin
# @Date     ：2023/3/28 13:14
# @Software : win10 python3.6
"""
采集中国共产党新闻网，习大大相关新闻、求是、原文
"""
import requests
from lxml import etree
import datetime
import os
import docx
import time
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from utils.CommonUtils import word_to_pdf


class CpcSpider(object):

    def __init__(self, current_year=datetime.datetime.today().year, current_month=datetime.datetime.today().month):
        self.current_year = current_year
        self.current_month = current_month
        self.num = 0
        self.readme = ''
        self.file_path = 'C:/Users/yang/Desktop/data'

    def start_request(self):
        url = 'http://cpc.people.com.cn/xijinping/'
        res = requests.get(url)
        html = etree.HTML(res.content.decode('gbk'))
        url_list = html.xpath('//h3[@class="white"]/span/a/@href')
        category_list = html.xpath('//h3[@class="white"]/span/a/text()')
        for i in range(len(url_list)):
            page = 1
            self.num = 0
            url = url_list[i].replace('index.html', f'index{page}.html')
            category = category_list[i]
            if category == '讲话':
                category = '原文'
            if category == '其他':
                category = '求是'
            print(category)
            self.parse(url, category, page)
            self.readme += f'    {category}  {self.num}条数据\n'
        file_name = f'{self.file_path}/{self.current_year}年{self.current_month}月习主席讲话、活动报道汇编/readme.txt'
        with open(file_name, 'w', encoding='utf-8') as f:
            f.write(f'采集地址：http://cpc.people.com.cn/xijinping/\n采集结果统计：\n')
            f.write(self.readme)
        with open(file_name.replace('data', 'data1'), 'w', encoding='utf-8') as f:
            f.write(f'采集地址：http://cpc.people.com.cn/xijinping/\n采集结果统计：\n')
            f.write(self.readme)

    def parse(self, url, category, page):
        file_path = f'{self.file_path}/{self.current_year}年{self.current_month}月习主席讲话、活动报道汇编/{category}'
        if not os.path.exists(file_path):
            os.makedirs(file_path)
        res = requests.get(url)
        html = etree.HTML(res.content.decode('gbk'))
        elem_list = html.xpath('//div[@class="fl"]/ul/li')
        for elem in elem_list:
            content_pub_time = elem.xpath('./i/text()')[0].replace('[', '').replace(']', '')
            content_pub_year = int(content_pub_time.split('年')[0])
            content_pub_month = int(content_pub_time.split('年')[-1].split('月')[0])
            if content_pub_year < self.current_year or content_pub_month < self.current_month:
                return
            content_url = 'http://cpc.people.com.cn' + elem.xpath('./a/@href')[0]
            content_title = elem.xpath('./a/text()')[0].replace('\n', '')
            print(content_title, content_url, content_pub_time)
            content_text = self.parse_content(content_url)
            file_name = f'{file_path}/{content_pub_time.replace("年", "").replace("月", "").replace("日", "")} {content_title}.txt'
            with open(file_name, 'w', encoding='utf-8') as f:
                f.write(content_title + '\n')
                f.write(f'({content_pub_time})' + '\n')
                f.write(content_text)
            self.convert_file(file_name)
            self.num += 1
        nxt_page = page + 1
        url = url.replace(f'index{page}.html', f'index{nxt_page}.html')
        self.parse(url, category, page)

    def parse_content(self, url):
        text = ''
        res = requests.get(url)
        html = etree.HTML(res.content.decode('gbk'))
        elem_list = html.xpath('//div[@class="show_text"]/p')
        for elem in elem_list:
            temp_text_list = elem.xpath('.//text()')
            for temp_text in temp_text_list:
                text = text + temp_text.strip().replace('\xa0', '') + '\n'
        return text

    def convert_file(self, file_name):
        with open(file_name, 'r', encoding='utf-8') as f:
            data_list = f.readlines()
        doc = docx.Document()

        p1 = doc.add_paragraph()
        text1 = p1.add_run(data_list[0].strip())
        text1.font.size = Pt(20)  # 设置字体大小
        text1.bold = True  # 设置字体是否加粗
        text1.font.name = 'Times New Roman'  # 设置西文字体
        text1.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
        p1.alignment = WD_TABLE_ALIGNMENT.CENTER  # 文字居中

        p2 = doc.add_paragraph()
        text2 = p2.add_run(data_list[1].strip())
        text2.font.size = Pt(10)
        text2.italic = True
        text2.font.name = 'Times New Roman'
        text2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p2.alignment = WD_TABLE_ALIGNMENT.CENTER

        p3 = doc.add_paragraph()
        content = ''
        for data in data_list[2:]:
            if data.strip():
                content = content + '        ' + data
        text3 = p3.add_run(content)
        text3.font.size = Pt(10)
        text3.font.name = 'Times New Roman'
        text3.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        txt_name = file_name.split('/')[-1]
        word_name = txt_name.replace('txt', 'docx')
        pdf_name = txt_name.replace('txt', 'pdf')

        file_path = file_name.replace(txt_name, '').replace('data', 'data1')
        word_path = file_path + 'doc'
        pdf_path = file_path + 'pdf'
        if not os.path.exists(word_path):
            os.makedirs(word_path)
        if not os.path.exists(pdf_path):
            os.makedirs(pdf_path)
        word_path = word_path + '/' + word_name
        pdf_path = pdf_path + '/' + pdf_name
        doc.save(word_path)
        word_to_pdf(word_path, pdf_path)

    def run(self):
        self.start_request()


if __name__ == '__main__':
    cpc = CpcSpider()
    cpc.run()
